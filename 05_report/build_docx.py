"""
Convert 05_report/report.md → 05_report/report.docx with the Aqualia
visual system applied.

Usage:  python 05_report/build_docx.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "report.md"
OUT = ROOT / "report.docx"

NAVY = RGBColor(0x00, 0x2F, 0x5F)
AQUA = RGBColor(0x5D, 0xB9, 0xD9)
DEEP = RGBColor(0x0B, 0x4F, 0x74)
SLATE = RGBColor(0x4A, 0x56, 0x63)
INK = RGBColor(0x2C, 0x33, 0x3B)   # near-black body/heading text
SAND_RGB = "d6cdb7"
GRID_RGB = "e6e8eb"

BODY_FONT = "Arial"  # Inter fallback — Arial is universal
HEADING_FONT = "Arial"


def set_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_paragraph_spacing(p, before=0, after=120, line=1.35) -> None:
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_styled_runs(p, text: str, base_font=BODY_FONT, base_size=11,
                    base_color=SLATE, bold=False) -> None:
    """Parse **bold** and `code` inline markdown within plain text.

    Bold runs stay in the body colour (slate). Colour is reserved for
    structural accents (headings, table headers, blockquote rail, cover
    page) so bolded emphasis in prose does not overload the eye with
    blue — a common machine-generated report tell.
    """
    if not text:
        return
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        r = p.add_run()
        r.font.name = base_font
        r.font.size = Pt(base_size)
        r.font.color.rgb = base_color
        if bold:
            r.bold = True
        if part.startswith("**") and part.endswith("**"):
            r.text = part[2:-2]
            r.bold = True
            # Keep slate. Do NOT colour inline bolds navy.
        elif part.startswith("`") and part.endswith("`"):
            r.text = part[1:-1]
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            r.text = part


def add_heading(doc: Document, text: str, level: int) -> None:
    """Custom heading styling with navy colour, bold, size scaled to level."""
    sizes = {1: 22, 2: 16, 3: 12, 4: 11}
    spacing_before = {1: 24, 2: 18, 3: 12, 4: 8}
    spacing_after = {1: 10, 2: 8, 3: 4, 4: 4}
    size = sizes.get(level, 11)

    # Level-1: add a thin aqua accent rectangle above
    if level == 1:
        accent = doc.add_paragraph()
        set_paragraph_spacing(accent, before=16, after=0, line=1.0)
        pPr = accent._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "24")
        bottom.set(qn("w:color"), "5DB9D9")
        pBdr.append(bottom)
        pPr.append(pBdr)
        # Make the paragraph empty but short so the border shows
        r = accent.add_run(" ")
        r.font.size = Pt(1)

    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=spacing_before[level],
                          after=spacing_after[level], line=1.2)
    r = p.add_run(text.strip())
    r.font.name = HEADING_FONT
    r.font.size = Pt(size)
    # Navy reserved for H1 and H2 (true section markers).
    # H3 and H4 use near-black ink so the page is not washed in blue.
    r.font.color.rgb = NAVY if level <= 2 else INK
    r.bold = True

    # Level-2 bottom rule
    if level == 2:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "5DB9D9")
        pBdr.append(bottom)
        pPr.append(pBdr)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=6, line=1.35)
    add_styled_runs(p, text)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=0, after=3, line=1.3)
    # Clear the default style's run formatting and re-add our styled runs.
    for r in list(p.runs):
        r.text = ""
    add_styled_runs(p, text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, before=0, after=3, line=1.3)
    for r in list(p.runs):
        r.text = ""
    add_styled_runs(p, text)


def add_blockquote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=6, line=1.35)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "5DB9D9")
    left.set(qn("w:space"), "6")
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent = Cm(0.4)
    # Quote body stays in slate. Aqua colour lives on the left rail
    # only, so there is at most one blue element per quote block.
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        r = p.add_run()
        r.font.name = BODY_FONT
        r.font.size = Pt(11)
        r.font.color.rgb = SLATE
        if part.startswith("**") and part.endswith("**"):
            r.text = part[2:-2]
            r.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            r.text = part[1:-1]
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r.text = part[1:-1]
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        else:
            r.text = part


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = tbl.rows[i].cells[j]
            set_cell_margins(cell)
            # Clear any default paragraph
            for p in list(cell.paragraphs):
                p.clear()
            cell_text = row[j] if j < len(row) else ""
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line=1.25)
            if i == 0:
                set_shading(cell, SAND_RGB)
                r = p.add_run(cell_text.replace("**", ""))
                r.font.name = BODY_FONT
                r.font.size = Pt(10)
                r.font.color.rgb = NAVY
                r.bold = True
            else:
                add_styled_runs(p, cell_text, base_size=10)


def parse_md(src_path: Path) -> list:
    """Return a list of block tuples: ('heading', level, text) |
    ('p', text) | ('bullet', text) | ('num', text) | ('blockquote', text) |
    ('table', [[cell,...],...]) | ('hr',) | ('blank',)."""
    blocks: list = []
    lines = src_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_front_matter = False
    in_list: str | None = None  # 'bullet' or 'num' or None

    def flush_table(pending_rows):
        # pending_rows: list of raw row strings like "| a | b |"
        cleaned = []
        for raw in pending_rows:
            raw = raw.strip()
            if raw.startswith("|"):
                raw = raw[1:]
            if raw.endswith("|"):
                raw = raw[:-1]
            cells = [c.strip() for c in raw.split("|")]
            cleaned.append(cells)
        # drop the separator row (---|---)
        cleaned = [r for r in cleaned if not all(
            re.fullmatch(r":?-{2,}:?", c or "") for c in r
        )]
        if cleaned:
            blocks.append(("table", cleaned))

    table_buffer: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Frontmatter guard
        if line.strip() == "---" and i == 0:
            in_front_matter = True
            i += 1
            continue
        if in_front_matter:
            if line.strip() == "---":
                in_front_matter = False
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            table_buffer.append(line)
            i += 1
            continue
        elif table_buffer:
            flush_table(table_buffer)
            table_buffer = []

        if not line.strip():
            blocks.append(("blank",))
            i += 1
            continue

        # Horizontal rule
        if re.fullmatch(r"\s*---+\s*", line):
            blocks.append(("hr",))
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            blocks.append(("heading", level, text))
            i += 1
            continue

        # Blockquote
        if line.startswith(">"):
            text = line[1:].strip()
            # merge multi-line blockquotes
            j = i + 1
            while j < len(lines) and lines[j].startswith(">"):
                text += " " + lines[j][1:].strip()
                j += 1
            blocks.append(("blockquote", text))
            i = j
            continue

        # Bulleted list
        m = re.match(r"^\s*[-*]\s+(.*)$", line)
        if m:
            blocks.append(("bullet", m.group(1).strip()))
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if m:
            blocks.append(("num", m.group(1).strip()))
            i += 1
            continue

        # Paragraph (merge continuation lines)
        text = line.strip()
        j = i + 1
        while (
            j < len(lines)
            and lines[j].strip()
            and not lines[j].strip().startswith(("#", ">", "|", "- ", "* "))
            and not re.match(r"^\s*\d+\.\s+", lines[j])
            and not re.fullmatch(r"\s*---+\s*", lines[j])
        ):
            text += " " + lines[j].strip()
            j += 1
        blocks.append(("p", text))
        i = j

    if table_buffer:
        flush_table(table_buffer)

    return blocks


def build(doc_path: Path, out_path: Path) -> None:
    blocks = parse_md(doc_path)

    doc = Document()

    # Page setup: 1" margins, US Letter. Inter-like body font.
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Set the default Normal style
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = SLATE

    # ─── Cover page ────────────────────────────────────────────────
    # Top kicker
    kicker = doc.add_paragraph()
    set_paragraph_spacing(kicker, before=72, after=8, line=1.0)
    r = kicker.add_run("IE SUSTAINABILITY DATATHON · MARCH 2026")
    r.font.name = BODY_FONT
    r.font.size = Pt(9)
    r.font.color.rgb = AQUA
    r.bold = True

    # Thin aqua rule under kicker
    rule = doc.add_paragraph()
    set_paragraph_spacing(rule, before=0, after=72, line=1.0)
    pPr = rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "5DB9D9")
    pBdr.append(bottom)
    pPr.append(pBdr)
    r = rule.add_run(" ")
    r.font.size = Pt(1)

    # Main title — very large, navy, bold
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, before=0, after=4, line=1.0)
    r = title.add_run("Water as")
    r.font.name = HEADING_FONT
    r.font.size = Pt(46)
    r.font.color.rgb = NAVY
    r.bold = True

    title2 = doc.add_paragraph()
    set_paragraph_spacing(title2, before=0, after=24, line=1.0)
    r = title2.add_run("Strategy.")
    r.font.name = HEADING_FONT
    r.font.size = Pt(46)
    r.font.color.rgb = AQUA
    r.bold = True

    # Subtitle
    sub = doc.add_paragraph()
    set_paragraph_spacing(sub, before=0, after=12, line=1.3)
    r = sub.add_run("A Double Materiality Framework for Aqualia, 2027–2030")
    r.font.name = HEADING_FONT
    r.font.size = Pt(15)
    r.font.color.rgb = DEEP
    r.bold = False

    # Lead paragraph on the cover
    lead = doc.add_paragraph()
    set_paragraph_spacing(lead, before=24, after=48, line=1.5)
    r = lead.add_run(
        "Three material topics. A €500 million EU-Taxonomy-aligned "
        "green bond programme. A 2027–2030 strategic roadmap with "
        "named owners and measurable KPIs. Simulated external ESG "
        "advisory engagement for Aqualia, the integrated water-cycle "
        "operator of the FCC Group."
    )
    r.font.name = BODY_FONT
    r.font.size = Pt(12)
    r.font.color.rgb = SLATE

    # Bottom band with metadata
    band = doc.add_paragraph()
    set_paragraph_spacing(band, before=120, after=6, line=1.0)
    pPr = band._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:color"), "002f5f")
    pBdr.append(top)
    pPr.append(pBdr)
    r = band.add_run(" ")
    r.font.size = Pt(1)

    meta = doc.add_paragraph()
    set_paragraph_spacing(meta, before=4, after=0, line=1.4)
    r = meta.add_run("Track   ")
    r.font.name = BODY_FONT
    r.font.size = Pt(8)
    r.font.color.rgb = AQUA
    r.bold = True
    r = meta.add_run("Double Materiality & ESG Strategy · Aqualia")
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY

    meta2 = doc.add_paragraph()
    set_paragraph_spacing(meta2, before=2, after=0, line=1.4)
    r = meta2.add_run("Team    ")
    r.font.name = BODY_FONT
    r.font.size = Pt(8)
    r.font.color.rgb = AQUA
    r.bold = True
    r = meta2.add_run("Los Gatos de Datos · IE Master in Business Analytics & Data Science")
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY

    meta3 = doc.add_paragraph()
    set_paragraph_spacing(meta3, before=2, after=0, line=1.4)
    r = meta3.add_run("Deck    ")
    r.font.name = BODY_FONT
    r.font.size = Pt(8)
    r.font.color.rgb = AQUA
    r.bold = True
    r = meta3.add_run("ma731.github.io/Aqualia-Datathon")
    r.font.name = BODY_FONT
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY

    # Page break after cover
    pb = doc.add_paragraph()
    pb_run = pb.add_run()
    pb_run.add_break(WD_BREAK.PAGE)

    # Render blocks — but skip the markdown's own title block (first H1 and its H2)
    # so we don't duplicate the title we just wrote.
    seen_h1 = False
    skip_title_h2 = False
    prev_was_blank = False

    for block in blocks:
        kind = block[0]
        if kind == "heading":
            level, text = block[1], block[2]
            # Suppress the title "Water as Strategy" and subtitle, since we have a title page.
            if level == 1 and not seen_h1:
                seen_h1 = True
                skip_title_h2 = True
                continue
            if level == 2 and skip_title_h2:
                skip_title_h2 = False
                continue
            skip_title_h2 = False
            add_heading(doc, text, level)
            prev_was_blank = False
        elif kind == "p":
            add_body_paragraph(doc, block[1])
            prev_was_blank = False
        elif kind == "bullet":
            add_bullet(doc, block[1])
            prev_was_blank = False
        elif kind == "num":
            add_numbered(doc, block[1])
            prev_was_blank = False
        elif kind == "blockquote":
            add_blockquote(doc, block[1])
            prev_was_blank = False
        elif kind == "table":
            add_table(doc, block[1])
            # add a tiny spacer after tables
            add_body_paragraph(doc, "")
            prev_was_blank = False
        elif kind == "hr":
            # Light grey divider — a fine hair between sections, not a
            # shouting navy underline. Colour matches our grid token.
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=8, after=10)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:color"), "c9ced3")
            pBdr.append(bottom)
            pPr.append(pBdr)
            prev_was_blank = False
        elif kind == "blank":
            # Skip consecutive blanks to avoid huge gaps
            if not prev_was_blank:
                # no-op; paragraph spacing handles gaps
                pass
            prev_was_blank = True

    # ─── Footer on every page ──────────────────────────────────────
    # Left: document title. Right: PAGE / NUMPAGES via Word field codes.
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.text = ""
        # Use a tab-separated layout so the title sits left and the
        # page number right. Word's default tabs for an A4/Letter page
        # with 1" margins put a right-tab stop at ~6.5".
        from docx.shared import Inches as _In
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(_In(6.5),
                               alignment=WD_ALIGN_PARAGRAPH.RIGHT)

        # Left side — document title in small caps
        r = p.add_run("Water as Strategy  ·  Aqualia Double Materiality  ·  Los Gatos de Datos")
        r.font.name = BODY_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = SLATE

        # Tab to right
        r = p.add_run("\t")
        r.font.name = BODY_FONT
        r.font.size = Pt(8)

        # Page X of Y via Word field codes
        def _add_field(paragraph, field_code):
            run = paragraph.add_run()
            fldChar1 = OxmlElement("w:fldChar")
            fldChar1.set(qn("w:fldCharType"), "begin")
            instrText = OxmlElement("w:instrText")
            instrText.set(qn("xml:space"), "preserve")
            instrText.text = f" {field_code} "
            fldChar2 = OxmlElement("w:fldChar")
            fldChar2.set(qn("w:fldCharType"), "end")
            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
            run.font.name = BODY_FONT
            run.font.size = Pt(8)
            run.font.color.rgb = NAVY
            run.bold = True
            return run

        _add_field(p, "PAGE")
        r = p.add_run(" / ")
        r.font.name = BODY_FONT
        r.font.size = Pt(8)
        r.font.color.rgb = SLATE
        _add_field(p, "NUMPAGES")

    doc.save(out_path)
    print(f"Wrote {out_path}  ({out_path.stat().st_size/1024:,.1f} KB)")


if __name__ == "__main__":
    build(SRC, OUT)
